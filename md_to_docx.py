"""
Convert bao cao Markdown sang DOCX.

Xu ly cac mau format thuc te trong "Bao cao Khai pha du lieu.md":
    **CHUONG X: ...**         -> Heading 1
    **X.Y. Tieu de muc**      -> Heading 2
    *X.Y.Z Tieu de muc con*   -> Heading 3 (in nghieng)
    **Hinh X.Y. ...**         -> Caption (in dam, can giua)
    ![alt](path)              -> Inline image
    - item                    -> Bullet list
    Doan thuong                -> Paragraph voi inline **bold** va *italic*
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parent
SRC = BASE_DIR / "Báo cáo Khai phá dữ liệu.md"
DST = BASE_DIR / "Báo cáo Khai phá dữ liệu.docx"


CHAPTER_RE = re.compile(r"^\*\*CHƯƠNG\s+\d+:.*\*\*$")
SECTION_RE = re.compile(r"^\*\*\d+(?:\.\d+)?\.\s.*\*\*$")
SUBSECTION_RE = re.compile(r"^\*\d+\.\d+\.\d+\.?\s.*\*$")
CAPTION_RE = re.compile(r"^\*\*Hình\s+\d+\.\d+\..*\*\*$")
IMG_RE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)$")
OTHER_BOLD_HEADING_RE = re.compile(r"^\*\*[^*]+\*\*$")
KETLUAN_RE = re.compile(r"^\*\*KẾT LUẬN\*\*$")
KL_SUB_RE = re.compile(r"^\*\d+\\?\.\s.*\*$")


def add_inline(paragraph, text: str) -> None:
    """Render inline **bold** va *italic* vao paragraph."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_image(doc: Document, alt: str, path_str: str) -> None:
    img_path = (BASE_DIR / path_str).resolve()
    if not img_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Thiếu hình: {path_str}]")
        run.italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(img_path), width=Cm(14))
    except Exception as e:
        p2 = doc.add_paragraph()
        p2.add_run(f"[Lỗi đọc hình {path_str}: {e}]").italic = True


def parse_blocks(md: str) -> list[list[str]]:
    """Tach noi dung Markdown thanh cac block (cach nhau bang dong trong)."""
    blocks: list[list[str]] = []
    current: list[str] = []
    for raw in md.splitlines():
        line = raw.rstrip()
        if line.strip() == "":
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(line)
    if current:
        blocks.append(current)
    return blocks


def render_block(doc: Document, lines: list[str]) -> None:
    joined = "\n".join(lines).strip()

    # Anh
    m = IMG_RE.match(joined)
    if m:
        add_image(doc, m.group("alt"), m.group("path"))
        return

    # Tieu de chuong
    if CHAPTER_RE.match(joined):
        text = joined.strip("*").strip()
        doc.add_heading(text, level=1)
        return

    # KET LUAN -> heading 1
    if KETLUAN_RE.match(joined):
        doc.add_heading("KẾT LUẬN", level=1)
        return

    # Hinh caption
    if CAPTION_RE.match(joined):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(joined.strip("*").strip())
        run.bold = True
        return

    # Section X.Y.
    if SECTION_RE.match(joined):
        text = joined.strip("*").strip()
        doc.add_heading(text, level=2)
        return

    # Subsection 7.3.1 trong dau *
    if SUBSECTION_RE.match(joined):
        text = joined.strip("*").strip()
        doc.add_heading(text, level=3)
        return

    # Tieu de phu phan ket luan: *1\. Ket qua dat duoc*
    if KL_SUB_RE.match(joined):
        cleaned = joined.strip("*").strip().replace("\\.", ".")
        doc.add_heading(cleaned, level=2)
        return

    # Bullet list (block toan dong bat dau bang -)
    if all(l.lstrip().startswith("- ") for l in lines):
        for l in lines:
            text = l.lstrip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, text)
        return

    # Cac heading in dam khac (1 dong, toan bo nam giua **)
    if len(lines) == 1 and OTHER_BOLD_HEADING_RE.match(joined):
        doc.add_heading(joined.strip("*").strip(), level=2)
        return

    # Doan van thuong - cac dong trong cung block noi voi nhau bang xuong dong mem
    p = doc.add_paragraph()
    for i, l in enumerate(lines):
        if i > 0:
            p.add_run("\n")
        add_inline(p, l)


def build_doc() -> None:
    md = SRC.read_text(encoding="utf-8")
    blocks = parse_blocks(md)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    for blk in blocks:
        render_block(doc, blk)

    doc.save(DST)
    print(f"Da xuat: {DST}")


if __name__ == "__main__":
    build_doc()
